import os
import pandas as pd
from docx import Document as DocxDocument
import fitz  # PyMuPDF
import json
from app.retriever import add_documents
from llama_index.schema import Document

DOCS_FOLDER = "./docs"
SUPPORTED_EXTENSIONS = [".csv", ".pdf", ".docx", ".txt", ".json"]
SCANNED_FILES_TRACKER = "./data/scanned_files.txt"


def extract_csv_as_chunks(path, chunk_size=30):
    try:
        df = pd.read_csv(path, encoding="utf-8", on_bad_lines="skip")
        df = df.dropna(how="all")
        texts = []
        for i in range(0, len(df), chunk_size):
            chunk = df.iloc[i:i+chunk_size]
            text_chunk = chunk.to_string(index=False)
            doc = Document(
                page_content=text_chunk,
                metadata={
                    "source": os.path.basename(path),
                    "dir": os.path.basename(os.path.dirname(path)),
                    "start_row": i,
                    "end_row": i + chunk_size
                }
            )
            texts.append(doc)
        return texts
    except Exception as e:
        print(f"Erreur lecture CSV {path} : {e}")
        return []


def extract_pdf(path):
    try:
        text = ""
        with fitz.open(path) as pdf:
            for page in pdf:
                text += page.get_text()
        return [Document(page_content=text, metadata={"source": os.path.basename(path), "dir": os.path.basename(os.path.dirname(path))})] if text.strip() else []
    except Exception as e:
        print(f"Erreur lecture PDF {path} : {e}")
        return []


def extract_docx(path):
    try:
        doc = DocxDocument(path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [Document(page_content=text, metadata={"source": os.path.basename(path), "dir": os.path.basename(os.path.dirname(path))})] if text.strip() else []
    except Exception as e:
        print(f"Erreur lecture DOCX {path} : {e}")
        return []


def extract_txt(path):
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": os.path.basename(path), "dir": os.path.basename(os.path.dirname(path))})] if text.strip() else []
    except Exception as e:
        print(f"Erreur lecture TXT {path} : {e}")
        return []


def extract_json(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            content = json.load(f)
        text = json.dumps(content, indent=2, ensure_ascii=False)
        return [Document(page_content=text, metadata={"source": os.path.basename(path), "dir": os.path.basename(os.path.dirname(path))})]
    except Exception as e:
        print(f"Erreur lecture JSON {path} : {e}")
        return []


def get_scanned_files():
    if os.path.exists(SCANNED_FILES_TRACKER):
        with open(SCANNED_FILES_TRACKER, "r") as f:
            return set(f.read().splitlines())
    return set()


def save_scanned_file(filename):
    with open(SCANNED_FILES_TRACKER, "a") as f:
        f.write(filename + "\n")


def scan_documents(folder: str = DOCS_FOLDER):
    documents = []
    scanned_files = get_scanned_files()
    for root, _, files in os.walk(folder):
        for fname in files:
            path = os.path.join(root, fname)
            rel_path = os.path.relpath(path, folder)
            if rel_path in scanned_files:
                print(f"⏩ {rel_path} déjà scanné, ignoré.")
                continue

            print(f"🔍 Scanning {rel_path}")
            ext = os.path.splitext(fname)[1].lower()
            if ext == ".csv":
                chunks = extract_csv_as_chunks(path)
            elif ext == ".pdf":
                chunks = extract_pdf(path)
            elif ext == ".docx":
                chunks = extract_docx(path)
            elif ext == ".txt":
                chunks = extract_txt(path)
            elif ext == ".json":
                chunks = extract_json(path)
            else:
                print(f"⛔ Format non supporté : {rel_path}")
                continue

            if chunks:
                documents.extend(chunks)
                save_scanned_file(rel_path)

    if documents:
        add_documents(documents)
        print(f"✅ {len(documents)} documents ajoutés à la base vectorielle.")
    else:
        print("Aucun nouveau document trouvé.")
